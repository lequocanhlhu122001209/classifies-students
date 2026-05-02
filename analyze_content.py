#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Script to analyze Word document and project to add missing content
"""
import os
import re
from pathlib import Path
from docx import Document

ROOT = Path(__file__).resolve().parents[0]
WORD_FILE = ROOT / "Xây dựng module phân loại trình độ sinh viên dựa trên điểm số và hành vi học tập.docx"

def count_words(text: str) -> int:
    return len(re.findall(r"\S+", text or ""))

def get_word_content():
    """Extract current Word document content"""
    if not WORD_FILE.exists():
        print(f"❌ File không tồn tại: {WORD_FILE}")
        return None
    
    doc = Document(str(WORD_FILE))
    content = []
    total_words = 0
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            words = count_words(text)
            total_words += words
            content.append(text[:100] + ("..." if len(text) > 100 else ""))
    
    return {
        "total_words": total_words,
        "total_paragraphs": len(doc.paragraphs),
        "content_preview": content[:30]  # First 30 paragraphs
    }

def analyze_project():
    """Analyze project structure to find missing topics"""
    project_info = {
        "files": [],
        "directories": [],
        "key_modules": []
    }
    
    # Scan src directory
    src_path = ROOT / "src"
    if src_path.exists():
        for file in src_path.glob("*.py"):
            project_info["files"].append(file.name)
    
    # Scan backend
    backend_path = ROOT / "backend"
    if backend_path.exists():
        for file in backend_path.glob("**/*.py"):
            project_info["files"].append(str(file.relative_to(ROOT)))
    
    # Scan scripts
    scripts_path = ROOT / "scripts"
    if scripts_path.exists():
        for file in scripts_path.glob("*.py"):
            project_info["files"].append(file.name)
    
    # Scan docs
    docs_path = ROOT / "docs"
    if docs_path.exists():
        for file in docs_path.glob("*.md"):
            project_info["files"].append(file.name)
    
    return project_info

def main():
    print("=" * 60)
    print("PHÂN TÍCH FILE WORD VÀ PROJECT")
    print("=" * 60)
    
    # Get Word content
    print("\n📄 PHÂN TÍCH FILE WORD:")
    word_info = get_word_content()
    if word_info:
        print(f"   - Số từ hiện tại: {word_info['total_words']}")
        print(f"   - Số đoạn văn: {word_info['total_paragraphs']}")
        print(f"   - Cần thêm: {max(0, 14000 - word_info['total_words'])} từ (để đủ 40 trang)")
        print(f"\n   📋 Nội dung preview (30 đoạn đầu):")
        for i, para in enumerate(word_info['content_preview'][:10], 1):
            print(f"      {i}. {para}")
    
    # Get project info
    print("\n\n📁 PHÂN TÍCH PROJECT:")
    project_info = analyze_project()
    print(f"   Các file code:")
    for f in sorted(set(project_info["files"]))[:20]:
        print(f"      - {f}")
    
    print(f"\n   Tổng: {len(set(project_info['files']))} file được tìm thấy")

if __name__ == "__main__":
    main()
