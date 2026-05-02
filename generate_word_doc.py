#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Generate comprehensive Word document from project files
"""
import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parents[0]
WORD_FILE = ROOT / "Xây dựng module phân loại trình độ sinh viên dựa trên điểm số và hành vi học tập.docx"

def read_file_safely(path):
    """Read a file safely"""
    try:
        if path.exists():
            with open(path, 'r', encoding='utf-8') as f:
                return f.read()
    except:
        pass
    return ""

def add_heading(doc, text, level=1):
    """Add a heading to document"""
    doc.add_heading(text, level=level)

def add_para(doc, text, bold=False, italic=False):
    """Add a paragraph"""
    p = doc.add_paragraph(text)
    if bold or italic:
        for run in p.runs:
            run.bold = bold
            run.italic = italic
    return p

def generate_comprehensive_doc():
    """Generate comprehensive document"""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    
    # TITLE
    title = doc.add_heading('XÂY DỰNG MODULE PHÂN LOẠI TRÌNH ĐỘ SINH VIÊN', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('Dựa Trên Điểm Số và Hành Vi Học Tập', level=2)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('')  # Spacer
    
    # TABLE OF CONTENTS
    add_heading(doc, 'MỤC LỤC', 1)
    toc_items = [
        '1. Giới Thiệu Chung',
        '2. Phân Tích Yêu Cầu Hệ Thống',
        '3. Kiến Trúc Hệ Thống',
        '4. Mô Hình Dữ Liệu',
        '5. Thiết Kế Cơ Sở Dữ Liệu',
        '6. Xử Lý Dữ Liệu và Chuẩn Hóa',
        '7. Thuật Toán Phân Loại',
        '8. Phát Hiện Bất Thường',
        '9. Hệ Thống Chấm Điểm Tích Hợp',
        '10. API Endpoints',
        '11. Giao Diện Người Dùng',
        '12. Kết Quả Và Đánh Giá',
        '13. Kế Hoạch Triển Khai',
        '14. Tài Liệu Kỹ Thuật',
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # SECTION 1: Introduction
    add_heading(doc, '1. Giới Thiệu Chung', 1)
    add_para(doc, '''Hệ thống phân loại trình độ sinh viên là một giải pháp toàn diện nhằm đánh giá và phân loại sinh viên dựa trên các chỉ số định lượng về điểm số, hành vi học tập và kỹ năng cụ thể theo từng môn học. Hệ thống kết hợp các kỹ thuật học máy hiện đại như K-Means clustering và K-Nearest Neighbors (KNN) để đưa ra những đánh giá chính xác, công bằng và khách quan.''')
    
    add_para(doc, '''Mục tiêu chính của hệ thống là:
• Cung cấp công cụ đánh giá sinh viên toàn diện và khách quan
• Phát hiện những trường hợp bất thường (gian lận, sai sót)
• Giúp giáo viên và cán bộ quản lý hiểu rõ hơn về năng lực sinh viên
• Hỗ trợ quyết định về hỗ trợ học tập và can thiệp kịp thời
• Cung cấp dữ liệu cho phân tích xu hướng học tập''')
    
    add_heading(doc, 'Các Tính Năng Chính', 2)
    features = [
        'Phân loại tự động 4 mức: Xuất sắc, Khá, Trung bình, Yếu',
        'Tích hợp 3 phương pháp chuẩn hóa dữ liệu: MinMax, Z-Score, Robust',
        'Phát hiện tự động các trường hợp gian lận hoặc bất thường',
        'Đánh giá chi tiết 16 kỹ năng lập trình (4 kỹ năng × 4 môn)',
        'Chấm điểm tích hợp từ 3 thành phần: Bài tập (30%), Giữa kỳ (30%), Cuối kỳ (40%)',
        'Đồng bộ dữ liệu với Supabase (cloud database)',
        'Giao diện dashboard thân thiện với biểu đồ và bảng thống kê',
        'API REST đầy đủ cho tích hợp với các hệ thống khác'
    ]
    for i, feature in enumerate(features, 1):
        add_para(doc, feature, italic=True)
    
    doc.add_page_break()
    
    # SECTION 2: Requirements
    add_heading(doc, '2. Phân Tích Yêu Cầu Hệ Thống', 1)
    
    add_heading(doc, 'Yêu Cầu Chức Năng', 2)
    add_para(doc, '''Hệ thống phải thực hiện các chức năng sau:''')
    
    requirements = {
        'Tiếp Nhận Dữ Liệu': [
            'Nhập dữ liệu sinh viên từ file CSV',
            'Đấu nối với SQL Server để lấy dữ liệu điểm',
            'Tích hợp với Supabase để lưu trữ dữ liệu cloud',
            'Hỗ trợ cật nhật dữ liệu dạng batch và real-time'
        ],
        'Xử Lý Dữ Liệu': [
            'Kiểm tra và làm sạch dữ liệu (data cleaning)',
            'Xử lý giá trị thiếu (missing values)',
            'Chuẩn hóa dữ liệu theo 3 phương pháp',
            'Trích xuất Features từ dữ liệu thô'
        ],
        'Phân Loại': [
            'Sử dụng K-Means để phân cụm sinh viên',
            'Huấn luyện mô hình KNN dựa trên nhãn K-Means',
            'Dự đoán phân loại cho sinh viên mới',
            'Lựa chọn phương pháp chuẩn hóa tối ưu'
        ],
        'Phát Hiện Bất Thường': [
            'Xác định các trường hợp điểm cao + thời gian làm bài ngắn',
            'Phát hiện điểm cao + vắng số giờ nhiều',
            'Cảnh báo nộp bài trễ quá mức cho phép',
            'Ghi lại lý do bất thường để theo dõi'
        ],
        'Báo Cáo & Thống Kê': [
            'Hiển thị danh sách sinh viên với phân loại',
            'Thống kê số lượng sinh viên theo từng cấp',
            'Biểu đồ phân bố điểm số',
            'Chi tiết kỹ năng từng sinh viên'
        ]
    }
    
    for category, reqs in requirements.items():
        add_heading(doc, category, 3)
        for req in reqs:
            add_para(doc, req)
    
    doc.add_page_break()
    
    # SECTION 3: Architecture
    add_heading(doc, '3. Kiến Trúc Hệ Thống', 1)
    
    add_heading(doc, 'Cấu Trúc Tổng Quát', 2)
    add_para(doc, '''Hệ thống được thiết kế theo mô hình 3 tầng (3-Tier Architecture):''')
    
    architecture = {
        'Tầng Dữ Liệu (Data Layer)': [
            'Supabase PostgreSQL: Lưu trữ chính',
            'SQLite: Cache cục bộ',
            'CSV Files: Import/Export dữ liệu',
            'Bảng: students, course_scores, skill_evaluations, classifications, v.v...'
        ],
        'Tầng Xử Lý (Processing Layer)': [
            'Chuẩn hóa dữ liệu (Data Normalization)',
            'Trích xuất đặc trưng (Feature Extraction)',
            'K-Means Clustering',
            'KNN Classification',
            'Anomaly Detection',
            'Integrated Scoring'
        ],
        'Tầng Bảng Diễn (Presentation Layer)': [
            'Backend Flask API (Python)',
            'Frontend Dashboard (HTML/CSS/JavaScript)',
            'REST API Endpoints',
            'WebSocket cho cập nhật real-time'
        ]
    }
    
    for layer, components in architecture.items():
        add_heading(doc, layer, 3)
        for comp in components:
            add_para(doc, comp)
    
    doc.add_page_break()
    
    # SECTION 4: Data Model
    add_heading(doc, '4. Mô Hình Dữ Liệu', 1)
    add_para(doc, '''Hệ thống sử dụng 7 bảng chính để lưu trữ và tổ chức dữ liệu sinh viên.''')
    
    add_heading(doc, 'Tổng Quan 7 Bảng', 2)
    tables_overview = [
        ('students', 'Thông tin cơ bản sinh viên (ID, tên, lớp, khoa, giới tính)'),
        ('student_csv_data', 'Dữ liệu hành vi: điểm TB, giữa kỳ, cuối kỳ, bài tập, vắng, nộp muộn'),
        ('course_scores', 'Điểm chi tiết 4 môn: NMLT, KTLT, CTDL, OOP'),
        ('skill_evaluations', 'Đánh giá 16 kỹ năng: 4 kỹ năng × 4 môn'),
        ('classifications', 'Kết quả phân loại từ K-Means, KNN, phương pháp chuẩn hóa'),
        ('integrated_scores', 'Điểm tích hợp từ bài tập (30%) + giữa kỳ (30%) + cuối kỳ (40%)'),
        ('exercise_details', 'Chi tiết bài tập từng sinh viên: thời gian, điểm, bất thường')
    ]
    
    for table_name, description in tables_overview:
        add_para(doc, f'{table_name}: {description}')
    
    doc.add_page_break()
    
    # SECTION 5: Database Design
    add_heading(doc, '5. Thiết Kế Cơ Sở Dữ Liệu', 1)
    
    add_heading(doc, 'Bảng Students (Sinh Viên)', 2)
    add_para(doc, '''Bảng này lưu trữ thông tin cơ bản của mỗi sinh viên:''')
    add_para(doc, '''• student_id (INTEGER, Primary Key): Mã sinh viên duy nhất (VD: 125001001)
• name (VARCHAR): Họ và tên sinh viên
• class (VARCHAR): Lớp (VD: 22CT111, 22CT112)
• khoa (VARCHAR): Khoa (VD: CNTT, KTDN)
• sex (VARCHAR): Giới tính (Nam/Nữ)
• created_at (TIMESTAMP): Ngày tạo record''')
    
    add_heading(doc, 'Bảng student_csv_data (Dữ Liệu Hành Vi)', 2)
    add_para(doc, '''Lưu các chỉ số hành vi học tập quan trọng:''')
    add_para(doc, '''• student_id (INTEGER, Foreign Key): Liên kết với students
• midterm_score (FLOAT): Điểm giữa kỳ (0-10)
• final_score (FLOAT): Điểm cuối kỳ (0-10)
• homework_score (FLOAT): Điểm bài tập (0-10)
• total_score (FLOAT): Điểm tổng kết (0-10)
• attendance_rate (FLOAT): Tỷ lệ tham gia (0-1)
• assignment_completion (FLOAT): Tỷ lệ hoàn thành bài tập (0-1)
• study_hours_per_week (INTEGER): Số giờ học/tuần
• late_submissions (INTEGER): Số lần nộp muộn
• behavior_score_100 (INTEGER): Điểm hành vi (0-100)''')
    
    add_heading(doc, 'Bảng course_scores (Điểm Môn Học)', 2)
    add_para(doc, '''Lưu điểm chi tiết của 4 môn học lập trình:''')
    add_para(doc, '''• student_id (INTEGER, Foreign Key)
• course_code (VARCHAR): Mã môn (NMLT, KTLT, CTDL, OOP)
• score (FLOAT): Điểm tổng môn (0-10)
• time_minutes (INTEGER): Thời gian làm bài (phút)
• midterm_score (FLOAT): Điểm giữa kỳ môn
• final_score (FLOAT): Điểm cuối kỳ môn

Bốn môn học:
1. NMLT - Nhập Môn Lập Trình: Biến, Điều khiển, Vòng lặp, Hàm
2. KTLT - Kỹ Thuật Lập Trình: Mảng, Con trỏ, Chuỗi, File I/O
3. CTDL - Cấu Trúc Dữ Liệu và Giải Thuật: Linked List, Stack/Queue, Cây, Hash
4. OOP - Lập Trình Hướng Đối Tượng: Class, Kế thừa, Đa hình, Đóng gói''')
    
    doc.add_page_break()
    
    # SECTION 6: Data Processing
    add_heading(doc, '6. Xử Lý Dữ Liệu và Chuẩn Hóa', 1)
    
    add_heading(doc, 'Quy Trình Xử Lý Dữ Liệu', 2)
    add_para(doc, '''Dữ liệu trải qua các bước xử lý sau:''')
    add_para(doc, '''1. TIẾP NHẬN DỮ LIỆU: Dữ liệu từ Supabase, SQL Server, hoặc CSV
2. KIỂM TRA CHẤT LƯỢNG: Phát hiện giá trị NULL, ngoài phạm vi, bản sao
3. RỰA SẠCH DỮ LIỆU: Xóa/sửa dữ liệu lỗi, điền giá trị thiếu
4. TRÍCH XUẤT ĐẶC TRƯNG: Tạo 12 features từ dữ liệu thô
5. CHUẨN HÓA: Áp dụng một trong 3 phương pháp
6. PHÂN LOẠI: Chạy K-Means rồi KNN''')
    
    add_heading(doc, 'Ba Phương Pháp Chuẩn Hóa', 2)
    add_para(doc, '''• MinMax Scaling: x' = (x - min) / (max - min) → Giá trị trong [0, 1]
  Ưu điểm: Giữ nguyên phân bố, dễ hiểu
  Nhược điểm: Nhạy cảm với outliers

• Z-Score Normalization: x' = (x - mean) / std → Giá trị tập trung quanh 0
  Ưu điểm: Phù hợp với phân bố chuẩn
  Nhược điểm: Có thể có giá trị rất lớn hoặc rất nhỏ

• Robust Scaling: x' = (x - median) / IQR → Ít bị ảnh hưởng bởi outliers
  Ưu điểm: Tốt nhất để xử lý outliers
  Nhược điểm: Phạm vi giá trị lớn hơn''')
    
    add_heading(doc, 'Các Đặc Trưng Được Trích Xuất (12 Features)', 2)
    add_para(doc, '''Hệ thống sử dụng 12 đặc trưng chính, chia làm 2 nhóm:

NHÓM ĐIỂM SỐ (50% trọng số):
1. avg_score: Điểm trung bình tất cả môn
2. midterm_avg: Điểm giữa kỳ trung bình
3. final_avg: Điểm cuối kỳ trung bình
4. homework_avg: Điểm bài tập trung bình

NHÓM HÀNH VI (50% trọng số):
5. attendance_rate: Tỷ lệ tham gia học
6. behavior_score: Điểm hành vi (0-100)
7. assignment_completion: Tỷ lệ hoàn thành bài tập
8. avg_study_hours: Giờ học/tuần trung bình
9. submission_timeliness: Tỷ lệ nộp đúng hạn
10. score_stability: Độ ổn định điểm (1 - coefficient_variation)
11. late_submission_ratio: Tỷ lệ nộp muộn
12. normalization_method: Phương pháp chuẩn hóa được sử dụng''')
    
    doc.add_page_break()
    
    # SECTION 7: Classification Algorithm
    add_heading(doc, '7. Thuật Toán Phân Loại', 1)
    
    add_heading(doc, 'K-Means Clustering', 2)
    add_para(doc, '''K-Means là thuật toán clustering không giám sát (unsupervised learning) được sử dụng để phân sinh viên thành 4 cụm:''')
    add_para(doc, '''Quy trình K-Means:
1. Khởi tạo 4 centroid ngẫu nhiên trong không gian 12 chiều
2. Gán mỗi sinh viên vào centroid gần nhất
3. Cập nhật vị trí centroid = trung bình các sinh viên trong cụm
4. Lặp lại 2-3 cho đến khi hội tụ (centroid không thay đổi)
5. Gán nhãn cụm dựa trên điểm tổng hợp trung bình

Bốn mức phân loại:
• Xuất Sắc (Excellence): Điểm ≥ 8.0
• Khá (Good): Điểm 7.0 - 8.0
• Trung Bình (Average): Điểm 5.0 - 7.0
• Yếu (Poor): Điểm < 5.0''')
    
    add_heading(doc, 'K-Nearest Neighbors (KNN)', 2)
    add_para(doc, '''KNN là thuật toán giám sát (supervised learning) được huấn luyện trên nhãn từ K-Means:''')
    add_para(doc, '''Quy trình KNN:
1. Huấn luyện: Dùng toàn bộ dữ liệu đã được K-Means gán nhãn
2. Dự đoán: Với sinh viên mới, tìm K sinh viên gần nhất trong tập huấn luyện
3. Bỏ phiếu: Nhãn là class xuất hiện nhiều nhất trong K lân cận
4. Khoảng cách được tính bằng Euclidean distance trong không gian 12 chiều

Tham số K: Thường chọn K = 5 hoặc K = 7 để bỏ phiếu''')
    
    add_heading(doc, 'Kết Quả Phân Loại Cuối Cùng', 2)
    add_para(doc, '''Kết quả cuối cùng kết hợp:
• Dự đoán từ K-Means (unsupervised)
• Dự đoán từ KNN (supervised)
• Phương pháp chuẩn hóa được sử dụng
• Các quy tắc phát hiện bất thường

Nước đi cuối cùng: Nếu phát hiện bất thường, hệ thống sẽ điều chỉnh kết quả (hạ xuống hoặc nâng lên) dựa trên mức độ bất thường.''')
    
    doc.add_page_break()
    
    # SECTION 8: Anomaly Detection
    add_heading(doc, '8. Phát Hiện Bất Thường', 1)
    add_para(doc, '''Hệ thống có 6 quy tắc phát hiện bất thường để cảnh báo các trường hợp nghi vấn gian lận hoặc sai sót trong dữ liệu:''')
    
    anomalies = [
        ('Điểm Cao + Thời Gian Ngắn (Mức độ: Nghiêm Trọng)', 
         'Nếu: Điểm ≥ 8.5 AND Thời gian làm bài < 5 giờ\nHành động: Hạ xuống mức "Yếu", cùng ghi lại lý do'),
        
        ('Điểm Cao + Vắng Nhiều (Mức độ: Nghiêm Trọng)',
         'Nếu: Điểm ≥ 8.0 AND Tỷ lệ tham gia < 50%\nHành động: Hạ xuống mức "Yếu"'),
        
        ('Nộp Muộn Thường Xuyên (Mức độ: Trung Bình)',
         'Nếu: Số lần nộp muộn ≥ 10 lần\nHành động: Hạ 2 bậc phân loại'),
        
        ('Điểm Miễn Cưỡng Đạt (Mức độ: Thấp)',
         'Nếu: Điểm 5.0 - 5.5 AND Hành vi < 60 điểm\nHành động: Ghi nhận cảnh báo, cần tư vấn'),
        
        ('Sự Thay Đổi Điểm Bất Thường (Mức độ: Trung Bình)',
         'Nếu: |Điểm cuối kỳ - Điểm giữa kỳ| > 3.0\nHành động: Kiểm tra lại dữ liệu, có thể là tích cực hoặc chệch hướng'),
        
        ('Độc Lập Quá Cao (Mức độ: Thấp)',
         'Nếu: Hoàn thành bài tập < 30% trong khi điểm ≥ 7.0\nHành động: Ghi nhận, cần xác minh')
    ]
    
    for anomaly_name, anomaly_desc in anomalies:
        add_heading(doc, anomaly_name, 3)
        add_para(doc, anomaly_desc)
    
    doc.add_page_break()
    
    # SECTION 9: Integrated Scoring
    add_heading(doc, '9. Hệ Thống Chấm Điểm Tích Hợp', 1)
    add_para(doc, '''Hệ thống tính điểm tích hợp từ 3 thành phần chính bằng công thức:''')
    add_para(doc, '''Điểm Tích Hợp = (Bài Tập × 30%) + (Giữa Kỳ × 30%) + (Cuối Kỳ × 40%)''', bold=True)
    
    add_heading(doc, 'Ý Nghĩa Của Tỷ Lệ Trọng Số', 2)
    add_para(doc, '''• 30% Bài Tập: Đánh giá quá trình, khả năng làm việc thực hành, sự kiên trì
• 30% Giữa Kỳ: Đánh giá kiến thức ở giữa kỳ, tạo động lực học ôn
• 40% Cuối Kỳ: Đánh giá kiến thức tổng quát toàn kỳ, là yếu tố quyết định

Tỷ số này cân bằng giữa quá trình học (60%) và kỳ thi chính thức (40%), phù hợp với phương pháp đánh giá liên tục hiện đại.''')
    
    add_heading(doc, 'Bảng Phân Loại Theo Điểm Tích Hợp', 2)
    add_para(doc, '''• Xuất Sắc (Excellence): Điểm Tích Hợp ≥ 8.0
• Khá (Good): 7.0 ≤ Điểm Tích Hợp < 8.0
• Trung Bình (Average): 5.0 ≤ Điểm Tích Hợp < 7.0
• Yếu (Poor): Điểm Tích Hợp < 5.0

Các bước tính toán:
1. Lấy từ bảng student_csv_data: homework_score, midterm_score, final_score
2. Áp dụng công thức trên để tính integrated_score
3. Lưu vào bảng integrated_scores
4. Sử dụng điểm này để phân loại cuối cùng''')
    
    doc.add_page_break()
    
    # SECTION 10: API Endpoints
    add_heading(doc, '10. API Endpoints', 1)
    add_para(doc, '''Hệ thống backend được xây dựng bằng Flask, cung cấp các API RESTful sau:''')
    
    add_heading(doc, 'Các Endpoint Chính', 2)
    
    endpoints = [
        ('GET', '/api/students', 
         'Trả về danh sách tất cả sinh viên với điểm tích hợp và phân loại',
         'Query params: class (lọc theo lớp), page (phân trang)'),
        
        ('GET', '/api/student/<student_id>',
         'Trả về chi tiết một sinh viên: thông tin cơ bản, điểm, kỹ năng, phân loại, bất thường',
         'Path param: student_id (mã sinh viên)'),
        
        ('GET', '/api/statistics',
         'Thống kê tổng quan: số sinh viên, phân bố theo từng mức, biểu đồ',
         'Query params: class (thống kê theo lớp cụ thể)'),
        
        ('GET', '/api/courses',
         'Danh sách 4 môn học và 16 kỹ năng',
         'Không có tham số'),
        
        ('POST', '/api/classify',
         'Chạy lại phân loại với phương pháp chuẩn hóa được chỉ định',
         'Body: {"normalization_method": "minmax|zscore|robust"}'),
        
        ('POST', '/api/sync-supabase',
         'Đồng bộ dữ liệu lên Supabase',
         'Body: {"table_name": "students", "operation": "update"}'),
    ]
    
    for method, endpoint, description, details in endpoints:
        add_heading(doc, f'{method} {endpoint}', 3)
        add_para(doc, f'Mô tả: {description}')
        add_para(doc, f'Chi tiết: {details}')
    
    doc.add_page_break()
    
    # SECTION 11: UI
    add_heading(doc, '11. Giao Diện Người Dùng', 1)
    add_para(doc, '''Giao diện frontend được xây dựng bằng HTML, CSS, JavaScript vanilla (không dùng framework phức tạp).''')
    
    add_heading(doc, 'Bố Cục Dashboard', 2)
    add_para(doc, '''1. HEADER: Logo, tiêu đề, nút làm mới dữ liệu
2. SIDEBAR: Menu điều hướng, tùy chọn lọc
3. MAIN CONTENT:
   - Phần thống kê: Số lượng sinh viên theo từng mức
   - Biểu đồ pie/bar: Phân bố phân loại
   - Bảng danh sách: Sinh viên + điểm + phân loại
4. MODALS:
   - Chi tiết sinh viên: Thông tin đầy đủ, kỹ năng từng môn
   - Cảnh báo bất thường: Danh sách các trường hợp cần chú ý''')
    
    add_heading(doc, 'Tính Năng Chính', 2)
    add_para(doc, '''• Lọc theo lớp, mức phân loại, có bất thường
• Sắp xếp theo tên, điểm, mức phân loại
• Tìm kiếm nhanh tên hoặc mã sinh viên
• Nhấp vào sinh viên để xem chi tiết
• Biểu đồ thống kê cập nhật khi dữ liệu thay đổi
• Export dữ liệu ra Excel
• Cập nhật dữ liệu thủ công cho từng sinh viên
• Chạy lại phân loại với phương pháp khác''')
    
    doc.add_page_break()
    
    # SECTION 12: Results and Evaluation
    add_heading(doc, '12. Kết Quả Và Đánh Giá', 1)
    
    add_heading(doc, 'Độ Chính Xác của Mô Hình', 2)
    add_para(doc, '''Hệ thống được đánh giá dựa trên:
• Độ chính xác (Accuracy): Tỷ lệ dự đoán đúng / tổng số dự đoán
• Precision: Trong các dự đoán dương tính, bao nhiêu là đúng
• Recall: Trong các trường hợp thực tế dương, bao nhiêu được tìm thấy
• F1-Score: Trung bình hài hòa của Precision và Recall

Lưu ý: Do là bài toán clustering không có "ground truth" chắc chắn, đánh giá chủ yếu dựa trên:
• Tính ổn định: Cùng dữ liệu, khác lần chạy có cho kết quả tương tự
• Khả năng giải thích: Có thể giải thích tại sao sinh viên được phân vào mức đó
• Hiệu suất thực tế: Giao viên đánh giá kết quả có hợp lý''')
    
    add_heading(doc, 'So Sánh 3 Phương Pháp Chuẩn Hóa', 2)
    add_para(doc, '''MinMax Scaling:
  + Giữ nguyên phân bố dữ liệu
  + Dễ hiểu, công thức đơn giản
  - Nhạy cảm với outliers
  
Z-Score Normalization:
  + Phù hợp với phân bố chuẩn
  + Ít bị ảnh hưởng outlier hơn MinMax
  - Có thể tạo giá trị rất lớn/nhỏ
  
Robust Scaling:
  + Tốt nhất để xử lý outliers
  + IQR không bị ảnh hưởng bởi extreme values
  - Phạm vi giá trị rộng lớn''')
    
    doc.add_page_break()
    
    # SECTION 13: Deployment
    add_heading(doc, '13. Kế Hoạch Triển Khai', 1)
    
    add_heading(doc, 'Giai Đoạn 1: Chuẩn Bị (Tuần 1-2)', 2)
    add_para(doc, '''√ Chuẩn bị dữ liệu sinh viên hiện tại
√ Kiểm tra chất lượng dữ liệu
√ Thiết lập Supabase account
√ Cấu hình SQL Server connection
√ Tạo bảng dữ liệu trong Supabase''')
    
    add_heading(doc, 'Giai Đoạn 2: Triển Khai Hệ Thống (Tuần 3-4)', 2)
    add_para(doc, '''√ Deploy backend Flask trên server
√ Deploy frontend trên web server
√ Chạy phân loại lần đầu với dữ liệu hiện tại
√ Kiểm tra kết quả và điều chỉnh
√ Tạo tài liệu hướng dẫn sử dụng
√ Đào tạo cán bộ sử dụng hệ thống''')
    
    add_heading(doc, 'Giai Đoạn 3: Yêu Cầu Bảo Trì (Tuần 5+)', 2)
    add_para(doc, '''√ Giám sát hệ thống hàng ngày
√ Cập nhật dữ liệu mới hàng tuần
√ Chạy lại phân loại sau khi cập nhật dữ liệu
√ Giải quyết sự cố phát sinh
√ Tối ưu hóa hiệu năng dựa trên phản hồi
√ Thêm tính năng mới theo yêu cầu''')
    
    doc.add_page_break()
    
    # SECTION 14: Technical Documentation
    add_heading(doc, '14. Tài Liệu Kỹ Thuật', 1)
    
    add_heading(doc, 'Cấu Trúc Thư Mục Project', 2)
    add_para(doc, '''classifies-students/
├── backend/                 # Backend Flask
│   ├── app.py              # Ứng dụng chính
│   ├── requirements.txt     # Dependencies
│   └── routes/             # Các API endpoint
│       ├── students.py
│       ├── classify.py
│       ├── ranking.py
│       ├── statistics.py
│       └── lazy_classifier.py
├── frontend/               # Frontend
│   └── index.html          # Dashboard
├── src/                    # Core modules
│   ├── student_classifier.py         # K-Means + KNN
│   ├── skill_evaluator.py            # Đánh giá kỹ năng
│   ├── integrated_scoring_system.py  # Tính điểm tích hợp
│   ├── supabase_sync.py              # Đồng bộ Supabase
│   └── knn_clustering_normalizer.py  # Chuẩn hóa dữ liệu
├── scripts/                # Các script tiện ích
│   ├── classify_new_students.py      # Phân loại lô
│   ├── student_clustering.py         # Phân tích cụm
│   └── analyze_changes.py            # Phân tích thay đổi
├── db/                     # Database
│   └── migrations/         # SQL scripts
├── docs/                   # Tài liệu
│   ├── DATA_MODEL.md
│   └── SYSTEM_DOCUMENTATION.md
└── requirements.txt        # Dependencies chính''')
    
    add_heading(doc, 'Cài Đặt Và Chạy Hệ Thống', 2)
    add_para(doc, '''1. Clone repository
   git clone <repo-url>
   cd classifies-students

2. Tạo virtual environment
   python -m venv .venv
   .venv\\Scripts\\activate   (Windows)
   source .venv/bin/activate (Linux/Mac)

3. Cài đặt dependencies
   pip install -r requirements.txt
   pip install -r backend/requirements.txt

4. Cấu hình .env file
   SUPABASE_URL=https://your-project.supabase.co
   SUPABASE_KEY=your-anon-key
   SQL_SERVER_CONNECTION=...

5. Chạy migrations (nếu có)
   python scripts/migrate_sql_old_to_new.py

6. Khởi động backend
   cd backend
   python app.py

7. Mở frontend (mở file frontend/index.html trong trình duyệt)
   hoặc serve bằng web server:
   python -m http.server 8000''')
    
    add_heading(doc, 'Dependencies Chính', 2)
    add_para(doc, '''Backend:
• Flask 2.3.0+: Web framework
• Flask-CORS 4.0.0+: Hỗ trợ CORS
• Scikit-Learn 1.3.0+: ML algorithms (K-Means, KNN)
• Pandas 2.0.0+: Data manipulation
• NumPy 1.24.0+: Numerical computing
• Supabase 2.0.0+: Cloud database client
• Python-dotenv 1.0.0+: Biến môi trường
• Requests 2.28.0+: HTTP requests

Frontend:
• Vanilla HTML5/CSS3/JavaScript (không cần build tool)
• Chart.js: Vẽ biểu đồ
• DataTables: Bảng dữ liệu tương tác''')
    
    add_heading(doc, 'Lệnh Hữu Ích', 2)
    add_para(doc, '''# Phân loại lại tất cả sinh viên
python scripts/classify_new_students.py

# Đồng bộ dữ liệu lên Supabase
python src/supabase_sync.py

# Phân tích cụm K-Means
python scripts/analysis/student_clustering.py

# Kiểm tra validator
python scripts/utils/validate_classifier.py

# Chạy test
pytest tests/

# Format code
black src/ backend/ scripts/

# Lint code
flake8 src/ backend/ scripts/''')
    
    doc.add_page_break()
    
    # CONCLUSION
    add_heading(doc, 'Kết Luận', 1)
    add_para(doc, '''Hệ thống phân loại trình độ sinh viên dựa trên điểm số và hành vi học tập là một giải pháp toàn diện, khoa học và hiệu quả. Bằng cách kết hợp các kỹ thuật học máy hiện đại (K-Means, KNN) với các phương pháp chuẩn hóa dữ liệu thích hợp, hệ thống cung cấp những đánh giá chính xác, công bằng và dễ hiểu cho các stakeholder.''')
    
    add_para(doc, '''Những lợi ích chính:
• Tự động hóa quá trình phân loại, tiết kiệm thời gian
• Phát hiện các trường hợp bất thường để can thiệp kịp thời
• Chuẩn hóa tiêu chuẩn đánh giá giữa các lớp và kỳ
• Cung cấp dữ liệu chi tiết cho phân tích xu hướng
• Hỗ trợ quyết định hành chính (học bổng, khen thưởng, can thiệp)''')
    
    add_para(doc, '''Hệ thống được thiết kế với khả năng mở rộng cao, dễ bảo trì, và có thể tích hợp vào các hệ thống quản lý học sinh hiện có. Với tài liệu đầy đủ và quy trình triển khai rõ ràng, hệ thống sẽ mang lại giá trị lâu dài cho nhà trường.''')
    
    # Save
    doc.save(str(WORD_FILE))
    
    # Count words
    new_doc = Document(str(WORD_FILE))
    total_words = 0
    for para in new_doc.paragraphs:
        words = len(para.text.split())
        total_words += words
    
    print(f"✅ Hoàn thành!")
    print(f"📄 File: {WORD_FILE.name}")
    print(f"📊 Số từ: ~{total_words}")
    print(f"📄 Số trang: ~{int(total_words / 250)} trang")

if __name__ == '__main__':
    generate_comprehensive_doc()
