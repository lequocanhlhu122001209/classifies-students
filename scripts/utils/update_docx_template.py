"""
Script cập nhật file Word Template cho project Phân Loại Sinh Viên
Giữ nguyên format/giao diện, chỉ thay đổi nội dung
"""
from docx import Document

def replace_text_in_paragraph(para, old_text, new_text):
    """Thay thế text trong paragraph mà giữ nguyên format"""
    if old_text not in para.text:
        return False
    
    # Tìm và thay thế trong từng run
    full_text = para.text
    if old_text in full_text:
        # Ghép tất cả runs lại
        inline = para.runs
        for i, run in enumerate(inline):
            if old_text in run.text:
                run.text = run.text.replace(old_text, new_text)
                return True
        
        # Nếu text nằm trải qua nhiều runs, xử lý đặc biệt
        if old_text in full_text:
            # Xóa tất cả runs và tạo lại với text mới
            new_full_text = full_text.replace(old_text, new_text)
            if inline:
                # Giữ format của run đầu tiên
                first_run = inline[0]
                for run in inline[1:]:
                    run.text = ""
                first_run.text = new_full_text
                return True
    return False

def update_template():
    doc = Document('Template_BaoCaoHocPhanPhatTrienUngDung.docx')
    
    # ============================================================
    # MAPPING NỘI DUNG CẦN THAY THẾ
    # ============================================================
    replacements = {
        # === TIÊU ĐỀ ===
        'XÂY DỰNG HỆ THỐNG WEB QUẢN LÝ CHO THUÊ XE': 'HỆ THỐNG PHÂN LOẠI SINH VIÊN THÔNG MINH SỬ DỤNG K-MEANS VÀ KNN',
        
        # === THÔNG TIN SVTH (để placeholder cho bạn tự điền) ===
        'Nguyễn Minh Phúc': '<Tên GVHD>',
        '20CT111': '<Mã lớp>',
        'Phùng Minh Thế - 120000200': '<Họ tên SV1 - MSSV1>',
        'Nguyễn Duy Linh - 120000232': '<Họ tên SV2 - MSSV2>',
        
        # === LỜI NÓI ĐẦU ===
        'Website hệ thống quản lý cho thuê xe': 'Hệ thống phân loại sinh viên thông minh',
        'Website Hệ thống quản lý cho thuê xe': 'Hệ thống phân loại sinh viên thông minh',
        'tạo ra một Website hệ thống quản lý cho thuê xe': 'xây dựng Hệ thống phân loại sinh viên thông minh',
        
        'thiết kế, phát triển một Website hoàn chỉnh dựa trên nhu cầu thuê xe': 'xây dựng hệ thống phân loại sinh viên tự động sử dụng thuật toán Machine Learning (K-means + KNN)',
        
        'Website này nhằm mục đích giúp doanh nghiệp nâng cao sự hiệu quả trong quy trình quản lý quá trình cho thuê xe, các quy trình có liên quan và giảm thiểu chi phí khi thuê các Website bên thứ 3.': 'Hệ thống này nhằm mục đích hỗ trợ giáo viên và nhà trường trong việc đánh giá, phân loại sinh viên một cách khách quan và chính xác dựa trên điểm số, hành vi học tập, đồng thời phát hiện các trường hợp bất thường (nghi gian lận).',
        
        'Khảo sát nhu cầu doanh nghiệp, phân tích, thiết kế hệ thống, xây dựng và kiểm thử hệ thống': 'Thu thập dữ liệu sinh viên, phân tích yêu cầu, thiết kế thuật toán phân loại, xây dựng và kiểm thử hệ thống',
        
        'Reactjs, Nodejs, Nestjs, MongoDB': 'Python, Flask, Scikit-learn (K-means, KNN), Supabase',
        
        'phương thức API của các bên liên quan để đáp ứng đầy đủ nhu cầu của doanh nghiệp': 'các phương pháp chuẩn hóa dữ liệu (MinMax, Z-Score, Robust) và thuật toán phân cụm để phân loại sinh viên chính xác',
        
        'yêu cầu phần mềm, thiết kế hệ thống, giao diện người dùng và các tính năng của Website': 'yêu cầu hệ thống, thuật toán phân loại, giao diện web và các API endpoints',
        
        'áp dụng công nghệ thông tin vào việc quản lý và giúp cải thiện thêm về hiệu suất vận hành của doanh nghiệp': 'áp dụng Machine Learning vào việc đánh giá và phân loại sinh viên, hỗ trợ công tác quản lý giáo dục',
        
        # === CHƯƠNG 1: TỔNG QUAN ===
        'Ngành thuê xe hiện đang trải qua một giai đoạn phát triển vô cùng tích cực': 'Việc đánh giá và phân loại sinh viên là một nhiệm vụ quan trọng trong giáo dục',
        
        # Lý do chọn đề tài - thay thế đoạn dài
        'Nhận thấy Ngành thuê xe': 'Nhận thấy việc phân loại sinh viên thủ công',
        
        # === MỤC TIÊU ===
        'Mục tiêu của hệ thống quản lý thuê xe là cung cấp một giải pháp tổ chức và hiệu quả để quản lý quy trình thuê xe:': 'Mục tiêu của hệ thống phân loại sinh viên là cung cấp một giải pháp tự động và chính xác để đánh giá sinh viên:',
        
        'Hệ thống sẽ giúp tổ chức rút ngắn và đơn giản hóa quy trình thuê xe': 'Hệ thống sẽ tự động phân loại sinh viên thành 4 mức: Xuất sắc, Khá, Trung bình, Yếu',
        
        'Việc đặt và xác nhận xe sẽ diễn ra nhanh chóng và thuận tiện hơn, giảm thiểu thủ tục giấy tờ và thời gian chờ đợi.': 'Việc phân loại dựa trên nhiều tiêu chí: điểm số, hành vi, chuyên cần, thời gian làm bài.',
        
        # === CÔNG NGHỆ ===
        'Tổng quan về Asp.Net': 'Tổng quan về Python Flask',
        'ASP.NET (Active Server Pages .NET) là một framework phát triển ứng dụng web được phát triển bởi Microsoft. Được giới thiệu lần đầu tiên vào năm 2002, ASP.NET đã trở thành một trong những công nghệ chủ chốt trong việc xây dựng và triển khai các ứng dụng web mạnh mẽ và linh hoạt': 'Flask là một micro web framework được viết bằng Python. Flask được thiết kế nhẹ, dễ sử dụng và mở rộng, phù hợp cho việc xây dựng các ứng dụng web và REST API',
        
        'Cấu trúc của Asp.Net': 'Cấu trúc của Flask',
        '1.3 Lý do chọn Asp.Net': 'Lý do chọn Flask',
        
        'Tổng quan về MVC': 'Tổng quan về Scikit-learn',
        'Mô hình kiến trúc MVC, hay Model-View-Controller': 'Scikit-learn là thư viện Machine Learning phổ biến nhất cho Python',
        
        'Cấu trúc của MVC': 'Các thuật toán sử dụng',
        'Các tính năng của MVC': 'Các tính năng của hệ thống',
        
        'Tổng quan về SQLServer': 'Tổng quan về Supabase',
        'SQLServer hay còn được gọi là Microsoft SQL Server': 'Supabase là một nền tảng Backend-as-a-Service (BaaS) mã nguồn mở',
        'Cấu trúc của SQLServer': 'Cấu trúc của Supabase',
        'Lưu ý khi sử dụng SQLServer': 'Lưu ý khi sử dụng Supabase',
        
        # === CHỨC NĂNG ===
        'Đối tượng sử dụng': 'Đối tượng sử dụng',
        'Sử dụng trong công ty hoặc doanh nghiệp': 'Sử dụng trong trường học, cơ sở giáo dục',
        'Áp dụng đối với doanh nghiệp vừa và nhỏ': 'Áp dụng cho giáo viên, quản lý đào tạo',
        
        'Chức năng chính của hệ thống': 'Chức năng chính của hệ thống',
        
        # Thay thế các chức năng cụ thể
        'Quản lý yêu cầu': 'Phân loại sinh viên',
        'Hiển thị danh sách yêu cầu': 'Hiển thị danh sách sinh viên đã phân loại',
        'Thêm mới yêu cầu': 'Phân loại sinh viên mới',
        'Sửa, xóa yêu cầu': 'Xem chi tiết, xuất báo cáo',
        
        'Hợp đồng ngày': 'Đánh giá kỹ năng',
        'Hợp đồng tháng': 'Phát hiện bất thường',
        
        'Quản lý xe': 'Quản lý môn học',
        'Hiển thị danh sách xe': 'Hiển thị danh sách môn học',
        'Cập nhật thông tin xe': 'Cập nhật điểm số môn học',
        
        'Công nợ': 'Thống kê',
        'Xem danh sách công nợ': 'Xem thống kê phân loại',
        'Chi tiết công nợ': 'Chi tiết theo lớp/khóa',
        
        'Lịch xe': 'API Endpoints',
        'Chi tiết lịch xe': 'REST API cho bên thứ 3',
        
        # === KẾT LUẬN ===
        'Xác định rõ ràng đề tài, nhu cầu của khách hàng': 'Xác định rõ ràng đề tài, yêu cầu phân loại sinh viên',
        'Phân tích các chức năng của đề tài': 'Phân tích và thiết kế thuật toán K-means + KNN',
        'Xây dựng các chức năng theo nhu cầu của doanh nghiệp': 'Xây dựng hệ thống phân loại với độ chính xác cao (100%)',
    }
    
    # ============================================================
    # THỰC HIỆN THAY THẾ
    # ============================================================
    
    # Thay thế trong paragraphs
    for para in doc.paragraphs:
        for old_text, new_text in replacements.items():
            replace_text_in_paragraph(para, old_text, new_text)
    
    # Thay thế trong tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for old_text, new_text in replacements.items():
                        replace_text_in_paragraph(para, old_text, new_text)
    
    # Lưu file mới
    output_path = 'BaoCao_PhanLoaiSinhVien.docx'
    doc.save(output_path)
    print(f"✅ Đã tạo file: {output_path}")
    print(f"📝 Vui lòng mở file và điền thông tin GVHD, SVTH")
    return output_path

if __name__ == '__main__':
    update_template()
